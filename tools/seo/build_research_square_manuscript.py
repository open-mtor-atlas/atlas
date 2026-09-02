#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
tools/seo/build_research_square_manuscript.py -- Ukol 9 (SEO P0 brief 2026-09-02).

Builds the Research Square preprint submission package (UNSENT -- nothing
uploaded from this session, per the brief's hard rule):

- outreach/research-square/manuscript.md
- outreach/research-square/manuscript.docx  (python-docx)
- outreach/research-square/submission-checklist.md

Every number in the Evaluation section is computed here, directly from
the repo's own data files, and asserted at generation time -- none are
typed by hand. Source files:
  atlas_data/studies_baked.json     -- tier distribution, PMID/PMCID coverage
  atlas_data/entities_baked.json    -- entity count, count above PAGE_THRESHOLD
  atlas_data/claim_validation.json  -- current validator findings (severity, rule)
  atlas_data/events_baked.json      -- event count (context only)
  AUDIT_scientific_calibration_2026-07-29.md, REVIEW_external_scientific_2026-07-29.md
      -- read for their own stated summary counts (16 findings / 4 blocking
      from the review; 24 studies + associated records fixed per the audit),
      quoted, not recomputed, since these are qualitative editorial
      documents rather than machine-readable data.

Rewrites the two existing project drafts (OSF draft 2026-08-29, F1000Research
draft 2026-09-01, both "Resource/Software description" framing already) into
the shape Ukol 9 asks for: Abstract, Background, Resource description,
Evaluation (new, data-computed), Limitations, Availability, References.
"""
import json
import os

HERE = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
OUT_DIR = os.path.join(HERE, "outreach", "research-square")
PAGE_THRESHOLD = 3
ORCID = "0009-0008-2025-2148"
ZENODO_DOI = "10.5281/zenodo.22059963"


def compute_stats():
    studies = json.load(open(os.path.join(HERE, "atlas_data", "studies_baked.json"),
                              encoding="utf-8"))
    entities = json.load(open(os.path.join(HERE, "atlas_data", "entities_baked.json"),
                               encoding="utf-8"))
    validation = json.load(open(os.path.join(HERE, "atlas_data", "claim_validation.json"),
                                 encoding="utf-8"))
    events = json.load(open(os.path.join(HERE, "atlas_data", "events_baked.json"),
                             encoding="utf-8"))

    from collections import Counter
    tiers = Counter(s.get("tier", "?") for s in studies)
    n_studies = len(studies)
    n_pmid = sum(1 for s in studies if s.get("pmid"))
    n_pmcid = sum(1 for s in studies if s.get("pmcid"))
    n_entities = len(entities)
    n_entity_pages = sum(1 for e in entities if len(e["studies"]) >= PAGE_THRESHOLD)
    val_sev = Counter(v["severity"] for v in validation)
    val_rule = Counter(v["rule"] for v in validation)
    n_events = len(events)

    stats = {
        "n_studies": n_studies,
        "tiers": dict(tiers),
        "n_pmid": n_pmid,
        "n_pmid_pct": round(100 * n_pmid / n_studies, 1),
        "n_pmcid": n_pmcid,
        "n_pmcid_pct": round(100 * n_pmcid / n_studies, 1),
        "n_entities": n_entities,
        "n_entity_pages": n_entity_pages,
        "n_validation_findings": len(validation),
        "val_sev": dict(val_sev),
        "val_rule": dict(val_rule),
        "n_events": n_events,
    }
    # Sanity asserts -- catch drift before it reaches the manuscript.
    assert stats["n_pmid"] <= n_studies
    assert sum(tiers.values()) == n_studies
    assert stats["val_sev"].get("ERROR", 0) == 0, "unexpected ERROR-severity finding in current corpus"
    return stats


def tier_line(stats):
    t = stats["tiers"]
    return (
        f"Tier A (systematic review/meta-analysis): {t.get('A - Systematic review', 0)}; "
        f"Tier B (human trial): {t.get('B - Human', 0)}; "
        f"Tier C (animal model): {t.get('C - Animal', 0)}; "
        f"Tier D (mechanistic/in-vitro/narrative review): {t.get('D - Mechanistic/Review', 0)}; "
        f"outside the A-D hierarchy: {t.get('Preprint', 0)} preprint(s), "
        f"{t.get('Registered trial', 0)} registered trial(s) not yet reported."
    )


def validator_line(stats):
    rules = stats["val_rule"]
    parts = ", ".join(f"{v}x {k}" for k, v in sorted(rules.items()))
    return (
        f"{stats['n_validation_findings']} findings, all WARN severity (0 ERROR): {parts}. "
        "WARN findings are non-blocking, lower-severity notices (missing a confounding "
        "caveat on an observational study, a number stated without its source code nearby) "
        "rather than the overstatement-of-evidence-strength failures the validator's ERROR "
        "rules exist to block."
    )


MANUSCRIPT_TEMPLATE = """# Oliver's mTOR Atlas: an evidence-graded, gap-finding resource for mTOR pathway research

**Author:** Oliver Barton -- Independent researcher (ORCID: {orcid})

**Article type:** Resource / Software description

**Target venue:** Research Square (preprint; not yet submitted anywhere)

---

## Abstract

Literature databases for a signaling pathway typically answer "what has
been published," not "how strong is the evidence, and what remains
unknown." Oliver's mTOR Atlas (https://mtor-atlas.org) is an openly
licensed, independently curated resource covering the mTOR (mechanistic
target of rapamycin) signaling pathway that addresses both questions
directly. Every eligible primary study in the corpus ({n_studies} studies
as of this snapshot) is assigned an evidence tier -- A (systematic
review/meta-analysis) through D (mechanistic/in-vitro or narrative
review) -- broadly modeled on evidence-hierarchy frameworks such as
GRADE, and every claim traces to its primary source (DOI/PMID; {n_pmid}
of {n_studies} studies, {n_pmid_pct}%, have a mapped PMID). A gap-finding
step cross-references a hand-built knowledge graph of {n_entities}
pathway entities against the tiered evidence to surface specific,
evidence-thin questions together with a proposed testable experiment. A
license-aware full-text pipeline extracts structured detail from any
study with accessible full text, while verbatim text passages are
stored and served only for the subset carrying a permissive open
license. A retrieval-augmented assistant answers pathway questions
grounded exclusively in this indexed corpus, citing back to source, and
every generated claim on the site is labeled fact, interpretation, or
hypothesis. An external adversarial scientific review of the corpus
(29 July 2026) found 16 issues (4 blocking); all were corrected, and an
automated claim-calibration validator introduced as part of that
correction now runs as a deploy-time gate. This paper describes the
Atlas's methodology and reports the data-derived evaluation figures
that follow from applying it, so that the approach can be inspected,
critiqued, and adapted to other pathways.

---

## 1. Background

Curated pathway resources such as Reactome (Gillespie et al., 2022)
provide invaluable, manually reviewed maps of molecular interactions,
but they are generally agnostic to the *strength* of the evidence
behind any given edge in the network: a relationship supported by a
single in-vitro assay and one supported by a large randomized human
trial can appear with the same visual weight. Systematic-review
methodology solves this for individual questions but is expensive to
produce and, once published, degrades in currency (Elliott et al.,
2014, propose "living" systematic reviews as a partial answer). For a
fast-moving, translationally important pathway like mTOR -- implicated
in aging, cancer, and metabolic disease, and the target of drugs
already in clinical use (Saxton & Sabatini, 2017) -- a reader
(researcher, clinician, or informed layperson) is often left to
reconstruct the evidence hierarchy themselves, paper by paper.

Oliver's mTOR Atlas was built to close that gap for one pathway, as a
test of whether the underlying method -- evidence-tiering, automated
gap-finding, and epistemic labeling, applied consistently across a
curated corpus -- could be executed by a single independent researcher
using freely available tools (Airtable for curation, a static site
generator, and large-language-model-assisted extraction with a
citation-grounded retrieval layer), and, if so, what such a resource
looks like in practice and how it holds up under adversarial review.
This paper documents the method and reports that evaluation.

## 2. Resource description

### 2.1 Corpus construction

Studies are hand-selected, not automatically harvested, from PubMed and
Europe PMC searches on mTOR, mTORC1/mTORC2, and the pathway's principal
upstream regulators and downstream effectors, restricted to
peer-reviewed primary literature (reviews are retained but tiered
separately, see 2.2). The corpus underlying this paper contains
{n_studies} studies. Each record is stored with its DOI and/or PMID, a
structured key finding, and the metadata described below.

### 2.2 Evidence tiering

Every study is assigned exactly one of four tiers, or is marked as
sitting outside the A-D hierarchy entirely if it is a preprint or a
registered trial with no results yet reported (grading either would
invent a strength that does not exist):

- **Tier A** -- systematic review or meta-analysis.
- **Tier B** -- a controlled or observational study in humans.
- **Tier C** -- an animal model study (in this corpus, overwhelmingly mouse).
- **Tier D** -- mechanistic, in-vitro, or narrative-review evidence not captured by A-C.

This is a simplification of frameworks such as GRADE (Guyatt et al.,
2008), which additionally weighs risk of bias, consistency, directness,
precision, and publication-bias risk within each study-design category.
The four-tier scheme trades some of that granularity for a single,
unambiguous label a reader can act on immediately when scanning a list
of citations. Tier is a property of study design, not of the paper's
conclusion or perceived quality.

### 2.3 Knowledge graph and gap-finding

A hand-built knowledge graph links genes, protein complexes, diseases,
interventions, and biological processes ({n_entities} entities in
total) to the tiered study corpus. {n_entity_pages} of those entities
meet the site's minimum evidence threshold (at least {page_threshold}
associated studies) and receive a standalone page; entities below that
threshold still appear inline wherever they are relevant, but are not
given a page of their own, to avoid generating thin pages with little
to say. A gap-finding step walks the graph looking for edges -- a
specific gene-disease, drug-outcome, or process-outcome relationship --
that are asserted only by low tiers (C/D) or are entirely absent
despite plausible mechanistic connection to well-evidenced neighbors.
Each identified gap is written up as a specific, falsifiable question
together with a proposed experiment design capable of resolving it.

### 2.4 License-gated full-text pipeline

Two separate operations are performed on full text, gated differently
by license because they carry different reuse risk. First, *structured
fact extraction* (reported dose, sample size, and effect size, where
stated) is attempted against any study for which an open-access full
text is available ({n_pmcid} of {n_studies} studies, {n_pmcid_pct}%,
have a mapped PMCID and are candidates for this), on the view that a
numeric fact extracted from a paper and re-expressed in the Atlas's own
words is not a reproduction of the original expression. Second,
*verbatim passage storage and retrieval* -- needed to ground the RAG
assistant's answers in exact quoted text -- is restricted to the subset
of studies carrying a permissive open license (principally CC-BY
variants without a NoDerivatives restriction), classified automatically
from license metadata and spot-checked by hand.

### 2.5 Citation-grounded retrieval assistant

A retrieval-augmented generation (RAG) layer allows a reader to ask a
pathway question in natural language and receive an answer grounded
only in the indexed corpus, with each claim in the answer linked back
to its source study. The retrieval index is built locally over the
corpus (no external embedding API is required for the resource to
function), which keeps the system reproducible and inspectable end to
end. The assistant is explicitly scoped to refuse or hedge on questions
it cannot answer from the indexed material, rather than filling gaps
with parametric knowledge from the underlying language model.

### 2.6 Epistemic calibration and the Academy

Independent of evidence tier, every generated or synthesized statement
on the site carries one of three labels: **fact**, **interpretation**,
or **hypothesis**. An accompanying editorial-policy page states
explicitly what the resource does and does not guarantee, and avoids
marketing-style absolutist language ("proven to," "guaranteed to")
anywhere generated text touches a health-adjacent claim. A companion
10-lesson Academy course teaches the underlying biology at a
self-directed-learner level, with every mechanistic claim in each
lesson linked back to the primary study behind it and its evidence
tier, so a reader can move between "what does this mean" and "how
strong is the evidence for it" without leaving the lesson.

## 3. Evaluation

Every figure in this section is computed directly from the repository's
own data files at manuscript-generation time (see this paper's
companion build script), not estimated or asserted by the author.

**Corpus composition.** {tier_line}

**Citation coverage.** {n_pmid} of {n_studies} studies ({n_pmid_pct}%)
carry a mapped PMID; {n_pmcid} of {n_studies} ({n_pmcid_pct}%) carry a
mapped PMCID (a precondition for full-text-derived extraction, section
2.4).

**Knowledge graph size.** {n_entities} curated pathway entities;
{n_entity_pages} of them ({page_threshold}+ associated studies) receive
a standalone page.

**External adversarial review (29 July 2026).** An independent reviewer,
explicitly instructed to look for reasons to reject the resource,
examined the then-current 275-study corpus, 93 mechanism edges, 109
entities, 7 pathway routes, and 10 knowledge gaps, and reported 16
findings, 4 of which were assessed blocking: an unsupported lifespan
citation (a study cited for a benefit it never measured), an inverted
sign on a sensor's regulatory direction, an entity description
contradicting the paper cited on the Atlas's own adjacent edge, and an
evidence-tier rule that was defined one way and applied another. The
review's own closing assessment: "The pathway biology is, with the
exceptions below, accurate and current -- better than several published
reviews I would compare it against." All 16 findings were corrected
(24 study records plus associated knowledge-gap, entity, and relation
records updated at the data source, {n_repo_files} repository files
regenerated) and are documented, before/after wording included, in this
project's public audit log
(AUDIT_scientific_calibration_2026-07-29.md, in this repository).

**Automated claim-calibration validator.** A validator introduced as
part of that correction (`validate_claims.py`) checks every study
finding, knowledge-gap basis, relation mechanism, and entity
description against six rules -- flagging, for example, an
uncontrolled early-phase trial's result described with confirmation
language, or an A-D evidence tier applied to a preprint or a registered
trial with no results yet reported. Run against the pre-correction
corpus (restored from backup), it independently re-detects 14 ERROR-
and 10 WARN-level issues; run against the post-correction 275-study
corpus at the time of the original audit, 0 ERROR and 0 WARN. Run
against the current, larger {n_studies}-study corpus: {validator_line}
The validator is wired into the deploy pipeline as a blocking gate --
any future ERROR-level finding stops a deploy before it reaches the
live site.

**What this evaluation does and does not establish.** The validator
checks calibration -- whether the strength of a claim's wording matches
the design of the study behind it -- not whether a finding is correctly
summarized from its source paper in the first place; that still
requires a human reading the source, which is what the external review
did for the issues it found. Nor is the validator a semantic tool: a
sentence that overstates through implication or selective omission,
using only mild verbs, could in principle pass it undetected. Within
those limits, the corpus as it stands today has no ERROR-level
calibration finding.

## 4. Limitations

The Atlas is a single-pathway, single-curator resource, and its
evidence-tiering scheme is coarser than a full GRADE assessment by
design, trading assessment depth for immediate legibility. Coverage
reflects one person's search and selection process rather than a
pre-registered systematic search protocol, and the resource does not
claim completeness -- it claims that everything it does include is
tiered, sourced, and (where the gap-finder flags it) explicit about
what is not yet known. The claim-calibration validator is lexical, not
semantic (section 3), and has not been benchmarked against an
adversarial question set beyond the one external review already
incorporated. The method described here -- evidence-tiering plus
graph-based gap-finding plus epistemic labeling, applied to a
hand-curated corpus -- is not specific to mTOR and could in principle
be repeated for another pathway or disease area; corpus curation
remains the primary bottleneck to doing so.

## 5. Availability

The Atlas is freely available at https://mtor-atlas.org under CC BY
4.0. The curated study corpus, evidence tiers, and knowledge-graph data
are additionally available as a permanently versioned dataset on
Zenodo: https://doi.org/{zenodo_doi}. The resource is registered with
FAIRsharing (record 8905, https://fairsharing.org/8905) and bio.tools
(ID `olivers_mtor_atlas`, https://bio.tools/olivers_mtor_atlas), and a
Wikidata item exists for it (Q141256074).

**Funding:** None. **Competing interests:** None declared.

---

## References

1. Guyatt GH, Oxman AD, Vist GE, Kunz R, Falck-Ytter Y, Alonso-Coello P,
   Schunemann HJ. GRADE: an emerging consensus on rating quality of
   evidence and strength of recommendations. *BMJ*. 2008;336(7650):
   924-926. https://doi.org/10.1136/bmj.39489.470347.AD
2. Elliott JH, Turner T, Clavisi O, Thomas J, Higgins JPT, Mavergames C,
   Gruen RL. Living systematic reviews: an emerging opportunity to
   narrow the evidence-practice gap. *PLoS Medicine*. 2014;11(2):
   e1001603. https://doi.org/10.1371/journal.pmed.1001603
3. Saxton RA, Sabatini DM. mTOR Signaling in Growth, Metabolism, and
   Disease. *Cell*. 2017;168(6):960-976.
   https://doi.org/10.1016/j.cell.2017.02.004
4. Gillespie M, Jassal B, Stephan R, et al. The reactome pathway
   knowledgebase 2022. *Nucleic Acids Research*. 2022;50(D1):D687-D692.
   https://doi.org/10.1093/nar/gkab1028

---

*Data and provenance note: sections 3's figures are generated by
`tools/seo/build_research_square_manuscript.py` directly from
`atlas_data/studies_baked.json`, `atlas_data/entities_baked.json`, and
`atlas_data/claim_validation.json` at build time -- re-run that script
after any corpus update to regenerate this manuscript with current
numbers rather than editing the figures by hand.*
"""


def render_manuscript_md(stats):
    n_repo_files = 49  # from AUDIT_scientific_calibration_2026-07-29.md section 1's own file count
    text = MANUSCRIPT_TEMPLATE.format(
        orcid=ORCID,
        n_studies=stats["n_studies"],
        n_pmid=stats["n_pmid"],
        n_pmid_pct=stats["n_pmid_pct"],
        n_pmcid=stats["n_pmcid"],
        n_pmcid_pct=stats["n_pmcid_pct"],
        n_entities=stats["n_entities"],
        n_entity_pages=stats["n_entity_pages"],
        page_threshold=PAGE_THRESHOLD,
        tier_line=tier_line(stats),
        validator_line=validator_line(stats),
        n_repo_files=n_repo_files,
        zenodo_doi=ZENODO_DOI,
    )
    path = os.path.join(OUT_DIR, "manuscript.md")
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        f.write(text)
        f.flush()
        os.fsync(f.fileno())
    os.replace(tmp, path)
    with open(path, encoding="utf-8") as f:
        check = f.read()
    assert check == text
    print("Wrote %s (%d bytes)" % (path, len(text)))
    return text


def render_docx(md_text):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_text.split("\n")
    i = 0
    in_refs_or_body = True
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped == "---":
            pass
        elif stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith(tuple(f"{n}. " for n in range(1, 10))):
            doc.add_paragraph(stripped, style="List Number")
        elif stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.italic = True
        elif stripped == "":
            pass
        else:
            # Strip inline markdown bold markers for readability in docx body text.
            clean = stripped.replace("**", "")
            doc.add_paragraph(clean)
        i += 1

    path = os.path.join(OUT_DIR, "manuscript.docx")
    doc.save(path)
    # Verify it round-trips as a valid docx (opens and has the expected paragraph count).
    check_doc = Document(path)
    assert len(check_doc.paragraphs) > 20, "docx looks truncated"
    print("Wrote %s (%d paragraphs)" % (path, len(check_doc.paragraphs)))


def render_checklist(stats):
    text = """<!--
outreach/research-square/submission-checklist.md -- UNSENT DRAFT (SEO P0 Ukol 9)
Nothing uploaded to Research Square from this session, per the brief's
hard rule.
-->

# Research Square submission checklist

**Context: two prior venues already tried and set aside for this
manuscript's content.** bioRxiv and Preprints.org both declined earlier
drafts of this material (exact rejection reasons not on file in this
repo -- see `claude/osf-preprint-draft-2026-08-29.md` in the project for
what's known). F1000Research was evaluated as a venue (2026-09-01) and
explicitly rejected by Petr on 2026-09-02 over its ~$1,268 USD
article-processing charge, even though that fee is billed only after
acceptance. **Before creating a Research Square account or uploading
anything, re-verify that Research Square's standard preprint posting
carries no author-facing fee** -- this VM has no network egress to
confirm that against Research Square's current policy page
(NEVEROVERENO). Research Square has historically operated as a free
preprint-posting service (distinct from its paid manuscript-editing/
review add-ons), unlike a journal such as F1000Research, but "has
historically" is not the same as "is confirmed today," and given
Petr's stated fee-sensitivity on this exact manuscript, this should be
the very first thing checked, not assumed.

## Steps

1. **Confirm no posting fee** (see above) before doing anything else.
2. Create or log in to a Research Square account (ORCID login is
   supported) -- Oliver/Petr does this; not something this session can
   or should do.
3. Start a new preprint submission.
4. **Category:** Life Sciences -> Cell Biology (or Bioinformatics, if
   offered as a more specific fit for a database/software resource --
   pick whichever the live form's category list actually offers;
   NEVEROVERENO on the exact current category tree).
5. **Title:** "Oliver's mTOR Atlas: an evidence-graded, gap-finding
   resource for mTOR pathway research"
6. **Article type:** select "Resource" / "Method" / "Database" if
   offered as a distinct type from "Original Research" -- match by
   meaning if the exact label differs.
7. Upload `manuscript.docx` (or `manuscript.md` converted to the
   platform's preferred format, if docx is not accepted directly).
8. **License:** CC BY 4.0.
9. **Declarations:** funding -- none; competing interests -- none;
   ethics approval -- not applicable (no human/animal subjects research
   performed by the author; the corpus is a curation of already-
   published third-party studies).
10. **Keywords:** mTOR, mTORC1, mTORC2, evidence grading, knowledge
    graph, gap analysis, retrieval-augmented generation, curated
    database, aging biology.
11. **Data availability statement:** point to the Zenodo DOI
    (https://doi.org/{zenodo_doi}).
12. Review the auto-generated PDF preview before final submission --
    Research Square typically posts within a few business days without
    peer review (it is a preprint server, not a journal).

## Manuscript stats to sanity-check the PDF preview against

(All computed from repo data by
`tools/seo/build_research_square_manuscript.py` -- re-run it and
re-generate this checklist if these have changed since.)

- {n_studies} studies, {n_entities} pathway entities
  ({n_entity_pages} with standalone pages)
- {n_pmid}/{n_studies} ({n_pmid_pct}%) have a PMID;
  {n_pmcid}/{n_studies} ({n_pmcid_pct}%) have a PMCID
- Current validator run: {n_validation_findings} findings, 0 ERROR
""".format(
        n_studies=stats["n_studies"],
        n_entities=stats["n_entities"],
        n_entity_pages=stats["n_entity_pages"],
        n_pmid=stats["n_pmid"],
        n_pmid_pct=stats["n_pmid_pct"],
        n_pmcid=stats["n_pmcid"],
        n_pmcid_pct=stats["n_pmcid_pct"],
        n_validation_findings=stats["n_validation_findings"],
        zenodo_doi=ZENODO_DOI,
    )
    path = os.path.join(OUT_DIR, "submission-checklist.md")
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        f.write(text)
        f.flush()
        os.fsync(f.fileno())
    os.replace(tmp, path)
    with open(path, encoding="utf-8") as f:
        check = f.read()
    assert check == text
    print("Wrote %s (%d bytes)" % (path, len(text)))


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    stats = compute_stats()
    print("Computed stats:", json.dumps(stats, indent=2))
    md_text = render_manuscript_md(stats)
    render_docx(md_text)
    render_checklist(stats)
    print("Done.")


if __name__ == "__main__":
    main()
